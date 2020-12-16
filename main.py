from docx import Document
from docx.shared import Inches

document  = Document()
# Profile Picture
document.add_picture('profile.png',
                     width=Inches(1.5)
                     )
#Basic Information Like Name, Phone Number and Email
name = input("What is your full Name?")
phone_number = input("What is your Phone Number?")
email = input("What is your Email?")

# Heading Title
document.add_heading(name,0)

document.add_paragraph(
    phone_number + ' | ' +email)

# About Me Section
document.add_heading('Objective',0)
about_me = input('Write your Objective? ')
document.add_paragraph(about_me)

# Work Experience Section
document.add_heading('Work Experience',0)
para = document.add_paragraph()

no_of_experiences = int(input("How many Employers' Experience do you want to add "))
# 4

for i in range(no_of_experiences):
    company = input("Enter your Company? ")
    from_date = input('From Date ')
    to_date = input('To Date ')
    para.add_run(company + ' ').bold = True
    para.add_run(from_date + '-'+ to_date + '\n').italic = True
    experience_details = input("Describe your Experience at " +company)
    para.add_run(experience_details)
    para.add_run('\n')

document.save('cv.docx')